from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH as align

document = Document()

datas = [
    {'사번': '11111', '이름': '길동', '부서': '경영지원팀', '교육여부': True},
    {'사번': '22222', '이름': '우치', '부서': '경영지원팀', '교육여부': False},
    {'사번': '33333', '이름': '선영', '부서': '개발팀', '교육여부': True},
    {'사번': '44444', '이름': '예진', '부서': '개발팀', '교육여부': False},
    {'사번': '55555', '이름': '지훈', '부서': '개발팀', '교육여부': True},
    {'사번': '66666', '이름': '송강', '부서': '클라우드팀', '교육여부': False},
    {'사번': '77777', '이름': '은우', '부서': '클라우드팀', '교육여부': True},
    {'사번': '88888', '이름': '빈이', '부서': '클라우드팀', '교육여부': False}
]

teams = {
    "경영지원팀": ["엑셀 기본", "엑셀로 작업하기"],
    "개발팀": ["파이썬 핵심", "파이썬 데이터 분석"],
    "클라우드팀": ["클라우드 아키텍처", "멀티 하이브리드 환경 이해"]
}

# 파일에 내용을 기록
# datas 들을 반복하며 하나씩 가져온다
for person in datas:
    # 교육여부가 True 인 사람의 이름 으로 파일을 생성하여 저장한다
    if person['교육여부']:
        document = Document()
        name = person['이름']
        document.add_heading('신입사원 교육 정보입니다.', level=0).alignment \
            = align.CENTER
        document.add_paragraph('안녕하세요 {0} 님, '.format(name))
        document.add_paragraph('쌤즈에 입사하신 것을 환영합니다.')
        document.add_paragraph('귀하의 부서에서는 아래 교육을 진행합니다.').alignment \
            = align.CENTER
        dept = person['부서']
        for team in teams[dept]:
            paragraph = document.add_paragraph()
            paragraph.add_run(team).bold = True

        document.add_paragraph("감사합니다. \n 주식회사 쌤즈 인사팀").alignment \
            = align.RIGHT
        document.save("신입사원교육_{0}.docx".format(person['이름']))
