from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH as align

document = Document()
# Title 텍스트로 표시, 0 은 제목, 1은 Heading1, 2는 Heading2 텍스트로 표시
document.add_heading('파이썬으로 워드 문서 만들기 예제~' , level=0)
document.add_paragraph('워드를 다루는 다양한 기능을 수행할 수 있어요^^').alignment = align.CENTER

datas = ["워드 목록 만들기1 ", "워드 목록 만들기2 ","워드 목록 만들기3"]

for data in datas:
    document.add_paragraph(data, style='List Bullet')

# 새 라인을 추가. 볼드체와 이탤릭체등 텍스트 처리
paragraph = document.add_paragraph()
paragraph.add_run(datas[0]).bold = True
paragraph.add_run(datas[1]).italic = True
paragraph.add_run(datas[2])

document.add_paragraph('파이썬 로고 이미지 입니다~').alignment = align.CENTER
document.add_picture('python.png')
document.save('result_word_write.docx')
