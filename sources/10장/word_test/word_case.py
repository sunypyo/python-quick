from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH as align

datas = [
    {'사번': '11111', '이름': '길동', '부서': '경영지원팀', '교육여부': True},
    {'사번': '22222', '이름': '우치', '부서': '경영지원팀', '교육여부': False},
    {'사번': '33333', '이름': '선영', '부서': '개발팀', '교육여부': True},
    {'사번': '44444', '이름': '예진', '부서': '개발팀', '교육여부': False},
    {'사번': '55555', '이름': '지훈', '부서': '개발팀', '교육여부': True},
    {'사번': '66666', '이름': '송강', '부서': '클라우드팀', '교육여부': False},
    {'사번': '77777', '이름': '은우', '부서': '클라우드팀', '교육여부': True},
    {'사번': '88888', '이름': '빈이', '부서': '클라우드팀', '교육여부': False}
]

teams = {
    "경영지원팀": ["엑셀 기본", "엑셀로 작업하기"],
    "개발팀": ["파이썬 퀵스타트", "파이썬 데이터 분석"],
    "클라우드팀": ["클라우드 아키텍처", "멀티 하이브리드 환경 이해"]
}

# datas 들을 반복하며 하나씩 가져온다
# TODO
# datas 리스트(사원 리스트) 에서 반복하면서 사원 정보(딕셔너리) 하나씩 꺼낸다
for person in datas:
    # 한 사원 정보(딕셔너리) 에서 교육여부가 True 인 사람을 선택한다
    if person['교육여부']:
        # 워드 문서 생성. Document 객체 생성
        document = Document()
        # 사원의 이름을 가져온다
        name = person['이름']
        # 워드 문서에 내용을 기록. 제목 추가. 가운데 정렬
        document.add_heading('신입 사원 교육 정보입니다.',
                             level=0).alignment = align.CENTER
        # 워드 문서에 단락 내용을 추가
        # 정렬 방식은 기본적으로 왼쪽 정렬
        document.add_paragraph('안녕하세요 {0} 님,'.format(name))
        document.add_paragraph('루비페이퍼에 입사하신 것을 환영합니다.')
        document.add_paragraph('귀하의 부서에서는 아래 교육을 진행합니다.'
                               ).alignment = align.CENTER

        # 사원이 속한 부서명을 가져온다
        dept = person['부서']
        # 부서명으로 teams(딕셔너리) 에서 key 가 부서인 value 를 가져온다
        # value 가 리스트이므로, 리스트 요소의 개수만큼 반복한다
        for team in teams[dept]:
            # 부서의 교육 내용을 가져와서 워드 문서에 추가한다
            # 스타일 정보를 가지고 있는 text 인 run 으로 추가한다. bold 로 추가
            paragraph = document.add_paragraph()
            paragraph.add_run(team).bold = True
            # 단순 텍스트 단락을 추가 할때는 아래 코드처럼
            # paragraph = document.add_paragraph(team)

        document.add_paragraph(
            '감사합니다. \n 루비페이퍼 인사팀').alignment = align.RIGHT

        # 파일을 저장. 파일 이름은 사원의 이름을 포함하여 저장
        document.save('신입사원교육_{0}.docx'.format(name))
