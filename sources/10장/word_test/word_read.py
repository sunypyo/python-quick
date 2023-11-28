import docx

# word_read_test.docx 파일로부터 Document 객체 생성
document = docx.Document('word_read_test.docx')
print("Document 객체:", document)
print("Document 의 제목: ", document.paragraphs[0].text)
# print(document.paragraphs[1].text)
# print(document.paragraphs[2].text)

# 반복문으로
for item in document.paragraphs:
    print(item.text)

print("=" * 30)

# 첫번째 table 을 가져온다
table = document.tables[0]
# table 의 행의 길이 len(table.rows) 만큼 반복하면서
# 각 행의 cell(열) 의 paragraphs 의 문자열을 출력
for i in range(len(table.rows)):
    print(table.rows[i].cells[0].paragraphs[0].text)
for i in range(len(table.rows)):
    print(table.rows[i].cells[1].paragraphs[0].text)
for i in range(len(table.rows)):
    print(table.rows[i].cells[2].paragraphs[0].text)

print("*" * 50)

# 중첩반복문
# table 이 여러개 있을때 여러개의 table 반복하기. table 의 개수만큼 반복
# table 을 하나씩 꺼내온다
for table in document.tables:
    # 하나의 table 내의 row 의 개수만큼 반복
    for i in range(len(table.rows)):
        print("=" * 30)
        # 하나의 row 안의 cell 의 개수만큼 반복
        for x in range(len(table.rows[i].cells)):
            # 하나의 cell 내부의 paragraphs의 개수만큼 반복
            for y in range(len(table.rows[i].cells[x].paragraphs)):
                # 각 paragraphs 의 text 를 가져와서 출력
                print(table.rows[i].cells[x].paragraphs[y].text)
print("*" * 50)

print("만든 사람:", document.core_properties.author)
print("생성 일자:", document.core_properties.created)